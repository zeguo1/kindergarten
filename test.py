# coding:utf-8
from zipfile import ZipFile
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from io import BytesIO
from PIL import Image
import time
from config import config

from urllib.parse import unquote
from urllib import parse

from bs4 import BeautifulSoup
import http.client, mimetypes, urllib, json, requests, os
import datetime

import xlrd
from xlrd import xldate_as_datetime

from http.cookiejar import LWPCookieJar

local_session = requests.Session()
local_session.cookies = LWPCookieJar(filename='wjxCookies.txt')

headers = config.get('wjx').get('headers')

def cookie_login():
    try:
        # 从文件中加载cookies(LWP格式)
        local_session.cookies.load(ignore_discard=True)
        print(local_session.cookies)
    except Exception:
        print("Cookies未能加载，使用密码登录")
        login()

    else:
        verify_login()

def verify_login():
    mine_response = local_session.get(url=config.get('wjx').get('mine_url'), headers=headers)
    soup = BeautifulSoup(mine_response.text, 'lxml')
    if soup.find_all(text=config.get('wjx').get('username')):
        print(soup.find_all(text=config.get('wjx').get('username')))
        return 'success';
    else:
        login()

def login():
    baseurl = config.get('wjx').get('login_url')
    basepage = local_session.get(url=baseurl, headers=headers)
    data = {}
    soup = BeautifulSoup(basepage.text, 'lxml')
    # a = soup.findAll("input", {"type": "hidden"})
    for i in soup.findAll("input"):
        if i.attrs['name'] == 'UserName':
            data[i.attrs['name']] = config.get('wjx').get('username')
        elif i.attrs['name'] == 'Password':
            data[i.attrs['name']] = config.get('wjx').get('password')
        elif i.attrs['name'] == 'hfUserName':
            data[i.attrs['name']] = ''
        elif i.attrs['name'] == 'RememberMe':
            data[i.attrs['name']] = 'on'
        else:
            data[i.attrs['name']] = i.attrs['value']
    login = local_session.post(url=basepage.url, headers=headers, data=data)
    local_session.cookies.save()
    print('登录成功')


# def url2Dict(url):
#     query = urllib.parse.urlparse(url).query
#     return dict([(k, v[0]) for k, v in urllib.parse.parse_qs(query).items()])
#
# def get_qiniu_image(url,studentname):
#     cookie_login()
#     img = 'static/uploads/'+studentname+'.jpg'  # 保存在本地的图片
#     print(url)
#     # basepage = local_session.get(url=url, headers=headers, verify=False)
#     # headers['Referer']=url
#     # print(basepage.encoding)
#     # soup = BeautifulSoup(basepage.text, 'lxml')
#     # new_url = "https://www.wjx.cn" + soup.find("a", {"id": "hrefDown"}).attrs['href']
#     # # unqu_url = unquote(new_url,encoding='utf-8')
#     # images = local_session.get(url=new_url, headers=headers, verify=False, )
#     aa = local_session.get(url= url ,headers=headers ,allow_redirects=False)
#     wjxurl = aa.next.url
#     # wjxurl1 = unquote(wjxurl, 'utf-8')
#     # wjxurl2 = urllib.parse.quote(wjxurl, 'utf-8')
#     # wjxhttp = urllib3.PoolManager()  # 创建PoolManager对象生成请求, 由该实例对象处理与线程池的连接以及线程安全的所有细节
#     # wjxresponse = wjxhttp.request('GET', url=wjxurl, headers=headers) # get方式请求
#     request = urllib.request.Request(url=wjxurl)
#     reponse = urllib.request.urlopen(request)
#     jpg_ima = Image.open(BytesIO(reponse.read()))  # 打开图片
#     jpg_ima = jpg_ima.convert('RGB')  # 去掉图片的A通道
#     jpg_ima.save(img, "JPEG")  # 保存新的图片
#     # local_session = requests.Session()
#     # baseurl = 'https://www.wjx.cn/login.aspx'
#     # loginpage = local_session.get(url=baseurl, headers = headers)
#     # data = {}
#     # soup = BeautifulSoup(loginpage.text, 'lxml')
#     # # a = soup.findAll("input", {"type": "hidden"})
#     # for i in soup.findAll("input"):
#     #     if i.attrs['name'] == 'UserName':
#     #         data[i.attrs['name']] = '18661636361'
#     #     elif i.attrs['name'] == 'Password':
#     #         data[i.attrs['name']] = 'savor,123'
#     #     elif i.attrs['name'] == 'hfUserName':
#     #         data[i.attrs['name']] = ''
#     #     elif i.attrs['name'] == 'RememberMe':
#     #         data[i.attrs['name']] = 'on'
#     #     else:
#     #         data[i.attrs['name']] = i.attrs['value']
#     # login = local_session.post(url=loginpage.url, headers=headers, data=data)
#     # img = 'static/uploads/'+studentname+'.jpg'  # 保存在本地的图片
#     # basepage = local_session.get(url=url, headers = headers)
#     # # data = {}
#     # soup = BeautifulSoup(basepage.text, 'lxml')
#     # images = local_session.get(url="https://www.wjx.cn/" + soup.find("a", {"id": "hrefDown"}).attrs['href'], headers=headers)
#     # jpg_ima = Image.open(BytesIO(images.content))  # 打开图片
#     # jpg_ima = jpg_ima.convert('RGB')  # 去掉图片的A通道
#     # jpg_ima.save(img, "JPEG")  # 保存新的图片
#     return img

class CourseExcle(object):
    def __init__(self, file, tableid):
        self.book = xlrd.open_workbook(file)
        self.table = self.book.sheets()[tableid]
        self.tableName = self.book.sheet_names()[tableid]
        self.rows = self.table.nrows
        self.cols = self.table.ncols
        self.time = time

    def readCourseInfo(self,filetime):
        all_student = config.get('classone').get('new_student_name')
        repeat_list = {}  # 重复提交名单
        for i in range(1, self.rows):
            # print (type(all_student))
            # print (datetime.datetime.strptime(self.table.row_values(i)[1],'%Y/%m/%d %H:%M:%S').date())
            # print (filetime)
            # print (datetime.datetime.strptime(self.table.row_values(i)[1],'%Y/%m/%d %H:%M:%S').date().__eq__(filetime))
            if datetime.datetime.strptime(self.table.row_values(i)[1],'%Y/%m/%d %H:%M:%S').date().__eq__(filetime):
                for j in all_student:
                    if all_student[j]['name'] == self.table.row_values(i)[6]:
                        if all_student[j]['status'] == 1:
                            repeat_list[all_student[j]['wjxid']] = self.table.row_values(i)[6]
                        print(j)
                        # print(all_student[j]['name']+self.table.row_values(i)[6])
                        all_student[j]['wjxid'] = int(self.table.row_values(i)[0])
                        all_student[j]['url'] = self.table.row_values(i)[7]
                        all_student[j]['status'] = 1
                        # print(i + 1)
                # data[i]=dict(zip(['name','url'],[self.table.row_values(i)[6],self.table.row_values(i)[7]]))
                        print(self.table.row_values(i)[6]+self.table.row_values(i)[7])
        return all_student,repeat_list


def download_image(url,name,false_list):
    # cookie_login()
    max_retry = 0
    while max_retry < 2:
        try:
            img = 'static/uploads/' + name + '.jpg'  # 保存在本地的图片
            # print(url)
            basepage = local_session.get(url=url, headers=headers)
            # data = {}
            soup = BeautifulSoup(basepage.text, 'lxml')
            images = local_session.get(url="https://www.wjx.cn/" + soup.find("a", {"id": "hrefDown"}).attrs['href'],headers=headers, allow_redirects=False)
            wjxurl = images.next.url
            request = urllib.request.Request(url=wjxurl)
            reponse = urllib.request.urlopen(request)
            jpg_ima = Image.open(BytesIO(reponse.read()))  # 打开图片
            jpg_ima = jpg_ima.convert('RGB')  # 去掉图片的A通道
            jpg_ima.save(img, "JPEG")  # 保存新的图片
            print(name+'获取成功')
            return img,false_list
        except Exception as e:
            print('错误 ：', e)
            print(name+'获取失败')
            false_list.append(name)
        max_retry += 1
    return img,false_list

def get_excle(year,month,day):
    cookie_login()
    exclepath = 'static/uploads/88015554_' + time.strftime("%Y%m%d%H%M%S", time.localtime()) + '.xlsx'
    baseurl = config.get('wjx').get('excle_url')
    resultpage = local_session.get(url=baseurl, headers=headers)
    s = open(exclepath,'wb')
    s.write(resultpage.content)
    s.close()
    excle = CourseExcle(exclepath, 0)
    result = excle.readCourseInfo(datetime.date(year,month,day))
    all_student=result[0]
    # 生成学生列表
    # 生成未提交学生列表
    student_list = {}
    undo_list = {}
    false_list = []
    for i in all_student:
        if all_student[i]['status'] == 1:
            student_list[all_student[i]['name']] = all_student[i]['wjxid']  # 给student_list元组添加学生信息
        elif all_student[i]['status'] == 0:
            undo_list[all_student[i]['name']] = i
    # get_images(exclepath)
    print(undo_list)
    doc = Document()
    doc.styles['Normal'].font.name = u'文星标宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'文星标宋')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    run = p.add_run('小一班学生及同住人健康码')
    run.font.size = Pt(18)
    for i in all_student:
        if all_student[i]['status'] == 1:
            # for i in reversed(unzipfile.namelist()):
            #     imgname = i.encode('cp437').decode('gbk')
            #     student_name = imgname.split('_')[1]  # 对文件名按照_切分,取出第2段
            #     if student_name in student_list:
            #         print(imgname)
            #         continue
            run = doc.add_paragraph().add_run(all_student[i]['name'])  # 添加文字
            # print (all_student[i]['name'])
            run.font.size = Pt(14)
            run = doc.add_paragraph().add_run()
            (img,false_list)=download_image(all_student[i]['url'], all_student[i]['name'],false_list)
            run.add_picture(img, width=Inches(3))
    print(false_list)
    downloadpath = 'static/download/'
    downloadfilename = '三明路小一班健康码' + time.strftime("%Y%m%d", time.localtime()) + '.docx'
    doc.save(downloadpath + downloadfilename)  # 保存路径

# def get_images(exclepath):
#     excle = CourseExcle(exclepath, 0)
#     result = excle.readCourseInfo(datetime.date(2020,9,27))
#     all_student=result[0]
#     # 生成学生列表
#     # 生成未提交学生列表
#     student_list = {}
#     undo_list = {}
#     for i in all_student:
#         if all_student[i]['status'] == 1:
#             student_list[all_student[i]['name']] = all_student[i]['wjxid']  # 给student_list元组添加学生信息
#         elif all_student[i]['status'] == 0:
#             undo_list[all_student[i]['name']] = i
#     # 生成爬取地址列表，保存图片，返回下载地址
#     for i in all_student:
#         if all_student[i]['status'] == 1:
#             # print (all_student[i]['name'])
#             download_image(all_student[i]['url'], all_student[i]['name'])
#             # done = executor.submit(download_image,all_student[i]['url'],session,all_student[i]['name'])
#             # done.add_done_callback(read_data)
#     print(undo_list)

get_excle(2020,10,12)