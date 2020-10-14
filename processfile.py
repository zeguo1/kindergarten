# coding:utf-8
from zipfile import ZipFile
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from PIL import Image
import time
from config import config

from bs4 import BeautifulSoup
import http.client, mimetypes, urllib, json, requests, os
import datetime

import xlrd
from xlrd import xldate_as_datetime

import threading
import urllib


# 1.从上传的压缩包中获取图片列表并获取人名及对应的文件名
# 2.打开word，根据列表，填写题目，人名字和图片

class CourseExcle(object):
    def __init__(self, file, tableid):
        self.book = xlrd.open_workbook(file)
        self.table = self.book.sheets()[tableid]
        self.tableName = self.book.sheet_names()[tableid]
        self.rows = self.table.nrows
        self.cols = self.table.ncols
        self.time = time

    def readCourseInfo(self,filetime):
        all_student = config.get('classone').get('new_student_name')
        repeat_list = {}  # 重复提交名单
        for i in range(1, self.rows):
            # print (type(all_student))
            # print (datetime.datetime.strptime(self.table.row_values(i)[1],'%Y/%m/%d %H:%M:%S'))
            # print (filetime)
            if datetime.datetime.strptime(self.table.row_values(i)[1],'%Y/%m/%d %H:%M:%S').date.__eq__(filetime):
                for j in all_student:
                    if all_student[j]['name'] == self.table.row_values(i)[6]:
                        if all_student[j]['status'] == 1:
                            repeat_list[all_student[j]['wjxid']] = self.table.row_values(i)[6]
                        print(j)
                        # print(all_student[j]['name']+self.table.row_values(i)[6])
                        all_student[j]['wjxid'] = int(self.table.row_values(i)[0])
                        all_student[j]['url'] = self.table.row_values(i)[7]
                        all_student[j]['status'] = 1
                        # print(i + 1)
                # data[i]=dict(zip(['name','url'],[self.table.row_values(i)[6],self.table.row_values(i)[7]]))
                        print(self.table.row_values(i)[6]+self.table.row_values(i)[7])
        return all_student,repeat_list

def get_word(filepath, filename):
    doc = Document()
    doc.styles['Normal'].font.name = u'文星标宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'文星标宋')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    run = p.add_run('小一班学生及同住人健康码')
    run.font.size = Pt(18)
    img = 'static/uploads/0.jpg'  # 保存在本地的图片
    unzipfile = ZipFile(filepath + filename, 'r')
    student_list = {}
    for i in reversed(unzipfile.namelist()):
        imgname = i.encode('cp437').decode('gbk')
        student_name = imgname.split('_')[1]  # 对文件名按照_切分,取出第2段
        if student_name in student_list:
            print(imgname)
            continue
        student_list[student_name] = imgname  # 给student_list元组添加学生信息
        run = doc.add_paragraph().add_run(student_name)  # 添加文字
        run.font.size = Pt(14)
        images = unzipfile.read(i)
        jpg_ima = Image.open(BytesIO(images))  # 打开图片
        jpg_ima = jpg_ima.convert('RGB')  # 去掉图片的A通道
        jpg_ima.save(img, "JPEG")  # 保存新的图片
        run = doc.add_paragraph().add_run()
        run.add_picture(img, width=Inches(3))
    print(student_list)
    downloadpath = 'static/download/'
    downloadfilename = '三明路小一班健康码' + time.strftime("%Y%m%d", time.localtime()) + '.docx'
    doc.save(downloadpath + downloadfilename)  # 保存路径
    return downloadfilename

def zip_get_result(filepath, filename):
    # 生成word文档标题
    doc = Document()
    doc.styles['Normal'].font.name = u'文星标宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'文星标宋')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    run = p.add_run('小一班学生及同住人健康码')
    run.font.size = Pt(18)
    img = 'static/uploads/0.jpg'  # 保存在本地的图片
    # 获取提交学生名单至student_list
    unzipfile = ZipFile(filepath + filename, 'r')
    all_student = config.get('classone').get('student_name')  # 全部学生名单
    student_list = {}  # 已提交名单
    repeat_list = {}  # 重复提交名单
    undo_list = {}  # 未提交名单
    for i in reversed(unzipfile.namelist()):
        imgname = i.encode('cp437').decode('gbk')
        student_filename = imgname.split('_')  # 对文件名按照_切分,取出第2段
        student_name = student_filename[1]
        all_student[student_name] = 1
        if student_name in student_list:
            repeat_list[student_filename[0]] = student_name
            # print(imgname)
            continue
        student_list[student_name] = student_filename[0]  # 给student_list元组添加学生信息
        run = doc.add_paragraph().add_run(student_name)  # 添加文字
        run.font.size = Pt(14)
        images = unzipfile.read(i)
        jpg_ima = Image.open(BytesIO(images))  # 打开图片
        jpg_ima = jpg_ima.convert('RGB')  # 去掉图片的A通道
        jpg_ima.save(img, "JPEG")  # 保存新的图片
        run = doc.add_paragraph().add_run()
        run.add_picture(img, width=Inches(3))
    # 保存word文档
    downloadpath = 'static/download/'
    downloadfilename = '三明路小一班健康码' + time.strftime("%Y%m%d", time.localtime()) + '.docx'
    doc.save(downloadpath + downloadfilename)  # 保存路径
    # 计算未提交人员名单
    for i in all_student:
        print(i, all_student[i])
        if all_student[i] == '0':
            undo_list[i] = all_student[i]
    result_list = {}  # 返回结果至页面
    result_list['student'] = student_list
    result_list['repeat'] = repeat_list
    result_list['undo'] = undo_list
    result_list['filename'] = downloadfilename
    result_list['student_size'] = len(student_list)
    result_list['repeat_size'] = len(repeat_list)
    result_list['undo_size'] = len(undo_list)
    return result_list

headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/84.0.4147.125 Safari/537.36',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
    'Accept-Encoding': 'gzip, deflate, br', 'Accept-Language': 'zh-CN,zh;q=0.9'}

headers_referer = {'Accept-Language': 'zh-cn', 'Accept': 'application/json, text/plain, */*',
                   'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 13_3_1 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148 MicroMessenger/7.0.10(0x17000a21) NetType/WIFI Language/zh_CN',
                   'Referer': 'http://kzyynew.qingdao.gov.cn:81/dist/index.html'}
headers_referer_origin = {'Accept-Language': 'zh-cn', 'Accept': 'application/json, text/plain, */*',
                          'Content-Type': 'application/json;charset=utf-8',
                          'User-Agent': 'Mozilla/5.0 (iPhone; CPU iPhone OS 13_3_1 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148 MicroMessenger/7.0.10(0x17000a21) NetType/WIFI Language/zh_CN',
                          'Referer': 'http://kzyynew.qingdao.gov.cn:81/dist/index.html',
                          'Origin': 'http://kzyynew.qingdao.gov.cn:81'}

# 该方法备用
def get_from_websearch():
    baseurl = 'https://www.wjx.cn/resultquery.aspx?activity=88015554'
    begintime = time.time()
    name='王麦荔'
    session = requests.Session()
    basepage = session.get(url=baseurl, headers=headers)
    print(type(basepage))
    va = urllib.parse.quote(('10000|' + name).encode('unicode-escape')).replace('%5Cu', '%u')
    print(type(basepage.text))
    querycond = '10000|' + name
    session.cookies['querycond88015554'] = urllib.parse.quote(('10000|' + name).encode('unicode-escape')).replace('%5Cu',
                                                                                                                  '%u')
    data = {}
    soup = BeautifulSoup(basepage.text, 'lxml')
    for i in soup.findAll("input", {"type": "hidden"}):
        if i.attrs['name'] == 'hfPostType':
            data[i.attrs['name']] = '1'
        elif i.attrs['name'] == 'hfQuery':
            data[i.attrs['name']] = '10000|' + name
        else:
            data[i.attrs['name']] = i.attrs['value']
    resultpage = session.post(url=baseurl, headers=headers, data=data)
    result = BeautifulSoup(resultpage.text, 'lxml')
    return result

def download_image(url,studentname):
    # if studentname == '宫沐暄':
    #     print('yes')
    img = 'static/uploads/'+studentname+'.jpg'  # 保存在本地的图片
    basepage = session.get(url=url, headers = headers)
    # data = {}
    soup = BeautifulSoup(basepage.text, 'lxml')
    # for i in soup.findAll("input"):
    #     if i.attrs['name'] == 'UserName':
    #         data[i.attrs['name']] = '18661636361'
    #     elif i.attrs['name'] == 'Password':
    #         data[i.attrs['name']] = 'savor,123'
    #     elif i.attrs['name'] == 'hfUserName':
    #         data[i.attrs['name']] = ''
    #     elif i.attrs['name'] == 'RememberMe':
    #         data[i.attrs['name']] = 'on'
    #     else:
    #         data[i.attrs['name']] = i.attrs['value']
    # login = session.post(url=basepage.url, headers=headers, data=data)
    # soup = BeautifulSoup(login.text, 'lxml')
    print()
    images = session.get(url="https://www.wjx.cn/" + soup.find("a", {"id": "hrefDown"}).attrs['href'], headers=headers)
    jpg_ima = Image.open(BytesIO(images.content))  # 打开图片
    jpg_ima = jpg_ima.convert('RGB')  # 去掉图片的A通道
    jpg_ima.save(img, "JPEG")  # 保存新的图片
    return img

def get_excle_from_web():
    exclepath = 'static/uploads/88015554_' + time.strftime("%Y%m%d%H%M%S", time.localtime()) + '.xlsx'
    global session
    session = requests.Session()
    baseurl = 'https://www.wjx.cn/login.aspx?returnurl=%2fwjx%2factivitystat%2fviewstatsummary.aspx%3factivity%3d88015554%26reportid%3d-1%26dw%3d1%26dt%3d2'
    basepage = session.get(url=baseurl, headers = headers)
    data = {}
    soup = BeautifulSoup(basepage.text, 'lxml')
    # a = soup.findAll("input", {"type": "hidden"})
    for i in soup.findAll("input"):
        if i.attrs['name'] == 'UserName':
            data[i.attrs['name']] = '18661636361'
        elif i.attrs['name'] == 'Password':
            data[i.attrs['name']] = 'savor,123'
        elif i.attrs['name'] == 'hfUserName':
            data[i.attrs['name']] = ''
        elif i.attrs['name'] == 'RememberMe':
            continue
        else:
            data[i.attrs['name']] = i.attrs['value']
    resultpage = session.post(url=baseurl, headers=headers, data=data)
    s = open(exclepath,'wb')
    s.write(resultpage.content)
    s.close()
    excle = CourseExcle(exclepath, 0)
    result = excle.readCourseInfo(datetime.date(2020,10,5))
    all_student=result[0]
    # 生成学生列表
    # 生成未提交学生列表
    student_list = {}
    undo_list = {}
    for i in all_student:
        if all_student[i]['status'] == 1:
            student_list[all_student[i]['name']] = all_student[i]['wjxid']  # 给student_list元组添加学生信息
        elif all_student[i]['status'] == 0:
            undo_list[all_student[i]['name']] = i
    # 生成爬取地址列表，保存图片，返回下载地址
    for i in all_student:
        if all_student[i]['status'] == 1:
            print (all_student[i]['name'])
            download_image(all_student[i]['url'],all_student[i]['name'])
    # 生成word
    doc = Document()
    doc.styles['Normal'].font.name = u'文星标宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'文星标宋')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    run = p.add_run('小一班学生及同住人健康码')
    run.font.size = Pt(18)
    student_list = {}
    for i in all_student:
        if all_student[i]['status'] == 1:
            # for i in reversed(unzipfile.namelist()):
            #     imgname = i.encode('cp437').decode('gbk')
            #     student_name = imgname.split('_')[1]  # 对文件名按照_切分,取出第2段
            #     if student_name in student_list:
            #         print(imgname)
            #         continue
            student_list[all_student[i]['name']] = all_student[i]['wjxid']  # 给student_list元组添加学生信息
            run = doc.add_paragraph().add_run(all_student[i]['name'])  # 添加文字
            print (all_student[i]['name'])
            run.font.size = Pt(14)
            download_image(all_student[i]['url'])
            run = doc.add_paragraph().add_run()
            run.add_picture(download_image(all_student[i]['url']), width=Inches(3))
        elif all_student[i]['status'] == 0:
            undo_list = {}
            undo_list[all_student[i]['wjxid']] = all_student[i]['name']

    print(student_list)
    downloadpath = 'static/download/'
    downloadfilename = '三明路小一班健康码' + time.strftime("%Y%m%d", time.localtime()) + '.docx'
    doc.save(downloadpath + downloadfilename)  # 保存路径

    result_list = {}  # 返回结果至页面
    result_list['student'] = student_list
    result_list['repeat'] = result[1]
    result_list['undo'] = undo_list
    result_list['filename'] = downloadfilename
    result_list['student_size'] = len(student_list)
    result_list['repeat_size'] = len(result[1])
    result_list['undo_size'] = len(undo_list)
    return result_list
# download_image('https://www.wjx.cn/wjx/viewfile.aspx?path=https%3a%2f%2fpubuserqiniu.paperol.cn%2f88015554_44_q2_1598243790PYmaWK.jpg%3fattname%3d44_2_pt2020_08_24_12_10_55.jpg&activity=88015554')
print (datetime.date(2020,10,5))
get_excle_from_web()
