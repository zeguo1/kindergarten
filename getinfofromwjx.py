# coding:utf-8
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from io import BytesIO
from PIL import Image
import time
import config

from bs4 import BeautifulSoup
import urllib, requests, os
import datetime

import xlrd
from xlrd import xldate_as_datetime

from http.cookiejar import LWPCookieJar

local_session = requests.Session()
local_session.cookies = LWPCookieJar(filename='wjxCookies.txt')

headers = config.config().readYaml().get('wjx').get('headers')

def cookie_login():
    try:
        # 从文件中加载cookies(LWP格式)
        local_session.cookies.load(ignore_discard=True)
        print(local_session.cookies)
    except Exception:
        print("Cookies未能加载，使用密码登录")
        login()
    else:
        verify_login()

def verify_login():
    mine_response = local_session.get(url=config.config().readYaml().get('wjx').get('mine_url'), headers=headers)
    soup = BeautifulSoup(mine_response.text, 'lxml')
    if soup.find_all(text=config.config().readYaml().get('wjx').get('username')):
        print(soup.find_all(text=config.config().readYaml().get('wjx').get('username')))
        print('登录成功')
        # return 'login success';
    else:
        login()

def login():
    baseurl = config.config().readYaml().get('wjx').get('login_url')
    basepage = local_session.get(url=baseurl, headers=headers)
    data = {}
    soup = BeautifulSoup(basepage.text, 'lxml')
    # a = soup.findAll("input", {"type": "hidden"})
    for i in soup.findAll("input"):
        if i.attrs['name'] == 'UserName':
            data[i.attrs['name']] = config.config().readYaml().get('wjx').get('username')
        elif i.attrs['name'] == 'Password':
            data[i.attrs['name']] = config.config().readYaml().get('wjx').get('password')
        elif i.attrs['name'] == 'hfUserName':
            data[i.attrs['name']] = ''
        elif i.attrs['name'] == 'RememberMe':
            data[i.attrs['name']] = 'on'
        else:
            data[i.attrs['name']] = i.attrs['value']
    login = local_session.post(url=basepage.url, headers=headers, data=data)
    soup = BeautifulSoup(login.text, 'lxml')
    if soup.find_all(text=config.config().readYaml().get('wjx').get('username')):
        local_session.cookies.save()
        print('用户名密码登录成功')
    else:
        print('登录失败'+soup.find(class_='submit-wrapper').span.text)

class CourseExcle(object):
    def __init__(self, file, tableid):
        self.book = xlrd.open_workbook(file)
        self.table = self.book.sheets()[tableid]
        self.tableName = self.book.sheet_names()[tableid]
        self.rows = self.table.nrows
        self.cols = self.table.ncols
        self.time = time

    def readCourseInfo(self,filetime):
        all_student =config.config().readYaml().get('classone').get('new_student_name')
        # print ('初始化列表：')
        # print (all_student)
        repeat_list = {}  # 重复提交名单
        for i in range(1, self.rows):
            # print (type(all_student))
            # print (datetime.datetime.strptime(self.table.row_values(i)[1],'%Y/%m/%d %H:%M:%S').date())
            # print (filetime)
            # print (datetime.datetime.strptime(self.table.row_values(i)[1],'%Y/%m/%d %H:%M:%S').date().__eq__(filetime))
            if datetime.datetime.strptime(self.table.row_values(i)[1],'%Y/%m/%d %H:%M:%S').date().__eq__(filetime):
                for j in all_student:
                    if all_student[j]['name'] == self.table.row_values(i)[6]:
                        if all_student[j]['status'] == 1:
                            repeat_list[all_student[j]['wjxid']] = self.table.row_values(i)[6]
                        print(j)
                        # print(all_student[j]['name']+self.table.row_values(i)[6])
                        all_student[j]['wjxid'] = int(self.table.row_values(i)[0])
                        all_student[j]['url'] = self.table.row_values(i)[7]
                        all_student[j]['status'] = 1
                        # print(i + 1)
                        # data[i]=dict(zip(['name','url'],[self.table.row_values(i)[6],self.table.row_values(i)[7]]))
                        print(self.table.row_values(i)[6]+self.table.row_values(i)[7])
        # print('加工完成列表：')
        # print(all_student)
        return all_student,repeat_list


def del_file(path_data):
    for i in os.listdir(path_data) :# os.listdir(path_data)#返回一个列表，里面是当前目录下面的所有东西的相对路径
        file_data = path_data + "\\" + i#当前文件夹的下面的所有东西的绝对路径
        if os.path.isfile(file_data) == True:#os.path.isfile判断是否为文件,如果是文件,就删除.如果是文件夹.递归给del_file.
            os.remove(file_data)
        else:
            del_file(file_data)

def download_image(url,name,false_list):
    # cookie_login()
    max_retry = 0
    while max_retry < 2:
        try:
            img = 'static/uploads/' + name + '.jpg'  # 保存在本地的图片
            # print(url)
            basepage = local_session.get(url=url, headers=headers)
            # data = {}
            soup = BeautifulSoup(basepage.text, 'lxml')
            images = local_session.get(url="https://www.wjx.cn/" + soup.find("a", {"id": "hrefDown"}).attrs['href'],headers=headers, allow_redirects=False)
            wjxurl = images.next.url
            request = urllib.request.Request(url=wjxurl)
            reponse = urllib.request.urlopen(request)
            jpg_ima = Image.open(BytesIO(reponse.read()))  # 打开图片
            jpg_ima = jpg_ima.convert('RGB')  # 去掉图片的A通道
            jpg_ima.save(img, "JPEG")  # 保存新的图片
            print(name+'获取成功')
            return img,false_list
        except Exception as e:
            print('错误 ：', e)
            print(name+'获取失败')
            false_list[name]=url
        max_retry += 1
    return img,false_list

def get_excle(resultdate_str):
    resultdate=datetime.datetime.strptime(resultdate_str, '%Y-%m-%d').date()
    # print(resultdate)
    cookie_login()
    exclepath = 'static/uploads/88015554_' + resultdate.isoformat() + '.xlsx'
    baseurl = config.config().readYaml().get('wjx').get('excle_url')
    resultpage = local_session.get(url=baseurl, headers=headers)
    s = open(exclepath,'wb')
    s.write(resultpage.content)
    s.close()
    excle = CourseExcle(exclepath, 0)
    result = excle.readCourseInfo(resultdate)
    all_student=result[0]
    repeat_list=result[1]
    # 生成学生列表
    # 生成未提交学生列表
    student_list = {}
    undo_list = {}
    for i in all_student:
        if all_student[i]['status'] == 1:
            student_list[all_student[i]['name']] = all_student[i]['wjxid']  # 给student_list元组添加学生信息
        elif all_student[i]['status'] == 0:
            undo_list[all_student[i]['name']] = i
    # get_images(exclepath)
    # print('未提交名单:')
    # print(undo_list)
    # print('提交名单:')
    # print(student_list)
    result_list = {}  # 返回结果至页面
    result_list['student'] = student_list
    result_list['repeat'] = repeat_list
    result_list['undo'] = undo_list
    result_list['student_size'] = len(student_list)
    result_list['repeat_size'] = len(repeat_list)
    result_list['undo_size'] = len(undo_list)
    if result_list['undo_size']==0:
        del_file('static/uploads')
        del_file('static/download')
        doc_info=generate_doc(resultdate,all_student)
        result_list['filename'] = doc_info[0]
        result_list['false'] = doc_info[1]
        result_list['false_size'] = len(doc_info[1])
    else:
        result_list['filename'] = ''
        result_list['false'] = ''
        result_list['false_size'] = -1
    # print('返回前端结果:')
    # print(result_list)
    return result_list

def generate_doc(resultdate,all_student):
    false_list = {}
    # 下载健康码并保存为word
    doc = Document()
    doc.styles['Normal'].font.name = u'文星标宋'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'文星标宋')
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中
    run = p.add_run('小一班学生及同住人健康码')
    run.font.size = Pt(18)
    for i in all_student:
        if all_student[i]['status'] == 1:
            run = doc.add_paragraph().add_run(all_student[i]['name'])  # 添加文字

            run.font.size = Pt(14)
            run = doc.add_paragraph().add_run()
            (img,false_list)=download_image(all_student[i]['url'], all_student[i]['name'],false_list)
            run.add_picture(img, width=Inches(3))
    print('失败列表：')
    print(false_list)
    downloadpath = 'static/download/'
    downloadfilename = '三明路小一班健康码' + resultdate.isoformat() + '.docx'
    doc.save(downloadpath + downloadfilename)  # 保存路径
    return downloadfilename,false_list
